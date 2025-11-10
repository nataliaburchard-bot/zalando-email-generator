import streamlit as st
import os
import requests
from docx import Document
from io import BytesIO

st.set_page_config(page_title="Zalando Email Generator", layout="centered")

st.title("📩 Gemini Email Generator")
st.markdown("Upload a `.doc` or `.docx` Jira ticket, and the app will generate the correct email.")

uploaded_file = st.file_uploader("Upload your Jira Word file", type=["doc", "docx"])
sender_name = st.text_input("Your name for sign-off (e.g., Natalia Burchard)")

def convert_doc_to_docx_cloudconvert(doc_file):
    api_key = os.getenv("CLOUDCONVERT_API_KEY")
    if not api_key:
        st.error("CloudConvert API key not set. Please configure it in the environment variables.")
        return None

  # 1. Request a signed import URL
signed_upload = requests.post(
    "https://api.cloudconvert.com/v2/import/upload",
    headers={"Authorization": f"Bearer {api_key}"}
).json()

upload_url = signed_upload["data"]["url"]
upload_params = signed_upload["data"]["parameters"]

# 2. Perform the actual upload using signed URL
files = {'file': (doc_file.name, doc_file, 'application/msword')}
upload_result = requests.post(upload_url, data=upload_params, files=files)

# 3. Use task name as input, not upload ID directly
import_task_name = signed_upload["data"]["id"]


    payload = {
        "tasks": {
            "import-upload": {
                "operation": "import/upload"
            },
            "convert-my-file": {
                "operation": "convert",
                "input": "import-upload",
                "input_format": "doc",
                "output_format": "docx"
            },
            "export": {
                "operation": "export/url",
                "input": "convert-my-file"
            }
        }
    }

    job_response = requests.post(
        "https://api.cloudconvert.com/v2/jobs",
        headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
        json=payload
    )

    job_id = job_response.json()["data"]["id"]

    # Polling the job status
    while True:
        status_response = requests.get(
            f"https://api.cloudconvert.com/v2/jobs/{job_id}",
            headers={"Authorization": f"Bearer {api_key}"}
        )
        status_data = status_response.json()
        if status_data["data"]["status"] == "finished":
            break

    export_task = next(task for task in status_data["data"]["tasks"] if task["name"] == "export")
    file_url = export_task["result"]["files"][0]["url"]

    # Download converted .docx
    docx_response = requests.get(file_url)
    return BytesIO(docx_response.content)

def extract_info_from_doc(doc):
    full_text = "\n".join([para.text for para in doc.paragraphs])

    # Email type detection
    if "quarantine storage" in full_text or "not ordered" in full_text.lower():
        email_type = "article_not_ordered"
    elif "price" in full_text.lower() and "invoice" in full_text.lower():
        email_type = "price_variance"
    else:
        email_type = "unknown"

    # Extract supplier name
    supplier = ""
    for para in doc.paragraphs:
        if "Supplier" in para.text:
            supplier = para.text.split(":")[-1].strip()
            break

    # Extract invoice number
    invoice = ""
    for para in doc.paragraphs:
        if "Supplier Invoice Number" in para.text:
            invoice = para.text.split(":")[-1].strip()
            break

    # Article table (if present)
    table_text = ""
    for table in doc.tables:
        for row in table.rows:
            table_text += "\t".join(cell.text.strip() for cell in row.cells) + "\n"

    return {
        "email_type": email_type,
        "supplier": supplier or "supplier",
        "invoice": invoice or "INVOICE-XXX",
        "table": table_text.strip()
    }

def generate_email(info, sender_name):
    if info["email_type"] == "price_variance":
        return f"""Dear {info["supplier"]},

I hope this email finds you well.
Your invoice {info["invoice"]} is currently in clarification due to a price discrepancy.

{info["table"]}

Please get back to us within the next 3 working days, otherwise we will have to deduct the difference automatically via debit note. 
If you have further questions, please do not hesitate to reach out.

Thank you and kind regards,  
{sender_name}"""
    elif info["email_type"] == "article_not_ordered":
        return f"""Dear {info["supplier"]},

I hope this email finds you well.
We have (an) item/s in quarantine storage as it looks like we have not ordered it and therefore we can not receive it/them. The articles have been delivered with [SN].

When items cannot be directly received due to specific issues they are sidelined and stored in our quarantine storage area (= a separate area in our warehouse). This additional clarification process is causing capacity losses and unforeseen costs.

{info["table"]}

We have two options:

1. Return (please provide the address, return label, and return authorization number).  
2. You can agree that we shall process the goods internally at Zalando's own discretion and thereby relinquish any and all rights that may have been reserved in the items (items will be processed further internally and then sold in bulk).

Please note, as per § 23 of the German Kreislaufwirtschaftsgesetz/Circular Economy Act (Kreislaufwirtschaftsgesetz) and similar legislation in other countries, we as a distributor are obliged to ensure that the quality of the articles we distribute is maintained and that they do not become waste (“duty of care”). 

In order to follow our sustainable and eco-friendly approach and by the German Circular Economy Act, we are unable to proceed with destruction and can offer a return or internal processing. Please note, although the law is a German one we do apply it to all our warehouses across Europe.

Please confirm how you would like to proceed within the next 3 working days.  
Should we not hear from you until then, we will assume your tacit consent that we may proceed by processing the articles internally.

Therefore, if you do not wish to accept this and would like to take back the articles for further processing on your side, please reach out to us within the deadline set.

If you have further questions, please do not hesitate to reach out.

Thank you and kind regards,  
{sender_name}"""
    else:
        return "⚠️ Could not determine the email type from this file. Please check the content or structure."

# Main app flow
if uploaded_file:
    file_name = uploaded_file.name
    if file_name.endswith(".doc"):
        st.info("Converting .doc to .docx using CloudConvert API...")
        docx_file = convert_doc_to_docx_cloudconvert(uploaded_file)
        if docx_file:
            doc = Document(docx_file)
        else:
            st.stop()
    else:
        doc = Document(uploaded_file)

    info = extract_info_from_doc(doc)
    email_output = generate_email(info, sender_name if sender_name else "[Your Name]")

    st.text_area("📨 Email Preview", email_output, height=600)
